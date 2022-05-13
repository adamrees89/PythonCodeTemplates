## Credit to Haider Imtiaz

# Automate MS WORD
# pip install python-docx
import docx
 
doc = docx.Document("zen_of_python.docx")
 
# read paragraph by paragraph
text = [p.text for p in doc.paragraphs]
print(text)
# read table by table
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            print(cell.text)
# Write in Doc
doc.add_paragraph("Hello World")
doc.save("test.docx")

#---------------------------------------------------

## Credit to Haider Imtiaz

#Convert DOCX to TXT
# First install this  "pip install docx2txt
import docx2txt
 
Text = docx2txt.process("filename.docx")
 
print(Text)

#---------------------------------------------------